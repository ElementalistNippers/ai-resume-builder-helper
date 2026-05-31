"""Generate a simple .docx resume from parsed data."""

from docx import Document
from docx.shared import Pt
from typing import Dict, List


def generate_docx(sections: Dict[str, List[str]], output_path: str) -> str:
    """Create a minimal .docx resume and save to output_path."""
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for heading, items in sections.items():
        if items:
            doc.add_heading(heading.capitalize(), level=2)
            for item in items:
                p = doc.add_paragraph(item)
                p.style = doc.styles["Normal"]

    doc.save(output_path)
    return output_path